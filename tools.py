import os
import docx
import fitz
import httpx
import chainlit as cl
from crawl4ai import AsyncWebCrawler, BrowserConfig, CrawlerRunConfig, CacheMode


async def web_search_tool(query: str, max_results: int = 3) -> list[dict]:
    """
    Performs a web search using Tavily's Developer API.
    Bypasses all datacenter blocks and returns optimized search results.
    """
    async with cl.Step(name="🔍 Searching via Tavily AI", type="tool") as step:
        step.input = query

        tavily_api_key = os.getenv("TAVILY_API_KEY")
        if not tavily_api_key:
            step.output = "Error: TAVILY_API_KEY is not set in the environment variables."
            return []

        url = "https://api.tavily.com/search"
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {tavily_api_key}"
        }
        payload = {
            "query": query,
            "max_results": max_results,
            "search_depth": "basic"
        }

        try:
            # Fully async POST call using httpx (already standard in our environment)
            async with httpx.AsyncClient() as client:
                response = await client.post(url, json=payload, headers=headers, timeout=15.0)

            if response.status_code == 200:
                data = response.json()
                results = data.get("results", [])

                found = []
                for r in results:
                    found.append({
                        "url": r.get("url"),
                        "title": r.get("title", "Source")
                    })

                if found:
                    urls_formatted = "\n".join([f"- [{r['title']}]({r['url']})" for r in found])
                    step.output = f"Discovered target URLs:\n{urls_formatted}"
                else:
                    step.output = "Tavily search returned no results."
                return found
            else:
                step.output = f"Tavily API returned status code {response.status_code}: {response.text}"
                return []

        except Exception as e:
            step.output = f"Tavily Search Error: {str(e)}"
            return []


async def scrape_source_tool(url: str, source_title: str) -> str:
    """
    Crawls a target page using Crawl4AI and displays dynamic progress in Chainlit.
    """
    domain = url.split("//")[-1].split("/")[0]
    
    async with cl.Step(name=f"🕷️ Crawling {domain}", type="tool") as step:
        step.input = url
        
        browser_cfg = BrowserConfig(headless=True, verbose=False)
        run_cfg = CrawlerRunConfig(cache_mode=CacheMode.BYPASS)
        
        async with AsyncWebCrawler(config=browser_cfg) as crawler:
            result = await crawler.arun(url=url, config=run_cfg)
            if result.success:
                word_count = len(result.markdown.split())
                step.output = f"Successfully scraped {word_count} words of text context."
                
                # FIX: Pass for_id=step.id so Chainlit binds the element correctly
                raw_source_element = cl.Text(
                    name=f"Source: {source_title}",
                    content=result.markdown,
                    display="side",
                    language="markdown"
                )
                await raw_source_element.send(for_id=step.id)
                
                return result.markdown
            else:
                step.output = f"Failed to crawl. Status: {result.status_code}"
                return ""


def extract_pdf_text(file_path: str) -> str:
    """
    Opens a PDF file from a local path and extracts plain text from all pages.
    """
    try:
        # Open the PDF using PyMuPDF
        doc = fitz.open(file_path)
        extracted_text = []

        # Iterate through pages and extract text
        for page_num, page in enumerate(doc, start=1):
            text = page.get_text()
            extracted_text.append(f"--- Page {page_num} ---\n{text}")

        doc.close()
        return "\n\n".join(extracted_text)
    except Exception as e:
        return f"Error reading PDF document: {str(e)}"

def extract_docx_text(file_path: str) -> str:
    """
    Opens a .docx file and extracts text from all paragraphs and tables.
    """
    try:
        doc = docx.Document(file_path)
        extracted_content = []
        
        # 1. Parse standard paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                extracted_content.append(para.text.strip())
                
        # 2. Parse tables (reconstructing cell rows with visual dividers)
        for table in doc.tables:
            for row in table.rows:
                # Compile cells that contain text
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    extracted_content.append(" | ".join(row_text))
                    
        return "\n\n".join(extracted_content)
    except Exception as e:
        return f"Error reading Word document: {str(e)}"